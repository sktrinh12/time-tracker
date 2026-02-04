from datetime import datetime
from docx import Document
from docx.shared import Inches
from config import (
    INVOICE_DIR,
    HOURLY_RATE,
    COMPANY_NAME,
    COMPANY_EMAIL,
    COMPANY_ADDR,
    CONSULTANT_ADDR,
    CONSULTANT_EMAIL,
    CONSULTANT_NAME,
    CONSULTANT_PHONE,
)
from db import entries_for_month, total_hours_for_month


def generate_monthly_invoice(month: str):
    """
    month: YYYY-MM
    """
    entries = entries_for_month(month)
    if not entries:
        raise ValueError(f"No entries found for this month - {month}")
    total_hours = total_hours_for_month(month)

    invoice_number = len(list(INVOICE_DIR.glob("Invoice_*"))) + 1
    filename = f"Invoice_{invoice_number:03}_{month}.docx"
    path = INVOICE_DIR / filename

    total_amount = total_hours * HOURLY_RATE
    today = datetime.today().date().isoformat()

    doc = Document()

    # ----------------------
    # Header
    # ----------------------
    doc.add_heading(CONSULTANT_NAME, level=1)
    doc.add_paragraph(CONSULTANT_ADDR)
    doc.add_paragraph(CONSULTANT_EMAIL)
    doc.add_paragraph(CONSULTANT_PHONE)

    doc.add_paragraph("")

    doc.add_paragraph(f"Invoice #: {invoice_number:03}")
    doc.add_paragraph(f"Invoice Date: {today}")
    doc.add_paragraph(f"Billing Month: {month}")

    doc.add_paragraph("")

    # ----------------------
    # Company
    # ----------------------
    doc.add_heading("Bill To:", level=2)
    doc.add_paragraph(COMPANY_NAME)
    doc.add_paragraph(COMPANY_ADDR)
    doc.add_paragraph(COMPANY_EMAIL)

    doc.add_paragraph("")

    # ----------------------
    # Table
    # ----------------------
    table = doc.add_table(rows=1, cols=5)
    hdr = table.rows[0].cells
    hdr[0].text = "Date"
    hdr[1].text = "Start"
    hdr[2].text = "End"
    hdr[3].text = "Hours"
    hdr[4].text = "Description"

    total_hours = 0.0

    for d, start, end, hours, desc in entries:
        row = table.add_row().cells
        row[0].text = d
        row[1].text = start
        row[2].text = end
        row[3].text = f"{hours:.2f}"
        row[4].text = desc or ""
        total_hours += hours

    doc.add_paragraph("")

    doc.add_paragraph(f"Hourly Rate: ${HOURLY_RATE:.2f}")
    doc.add_paragraph(f"Total Hours: {total_hours:.2f}")
    doc.add_paragraph(f"Total Amount Due: ${total_amount:,.2f}")

    doc.add_paragraph("")
    doc.add_paragraph("Thank you for your business.")

    doc.save(path)
    return path


def next_invoice_number():
    return len(list(INVOICE_DIR.glob("Invoice_*.txt"))) + 1
